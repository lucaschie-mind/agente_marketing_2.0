"""
Ingestão de materiais do Google Drive:
lê arquivos de todas as subpastas, extrai texto, faz chunking
e retorna documentos prontos para indexar no pgvector.

Estrutura esperada no Drive:
  📁 Raiz (GOOGLE_DRIVE_FOLDER_ID)
    📁 brand/       → brand book, tom de voz
    📁 modulos/     → one-pagers de cada módulo
    📁 personas/    → ICPs
    📁 cases/       → cases aprovados
    📁 conteudos/   → posts, blog, e-books
    📁 urls/        → arquivo urls.txt com URLs para raspar
"""

import hashlib
import io
import json
from pathlib import Path
from typing import Iterator

import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from src import config


# ── Google Drive ──────────────────────────────────────────────────────────────

def get_drive_service():
    from googleapiclient.discovery import build
    from google.oauth2.service_account import Credentials

    info = json.loads(config.GOOGLE_SERVICE_ACCOUNT_JSON)
    creds = Credentials.from_service_account_info(
        info, scopes=["https://www.googleapis.com/auth/drive.readonly"]
    )
    return build("drive", "v3", credentials=creds)


def list_drive_files(service, folder_id: str, folder_name: str = "") -> list[dict]:
    """Lista recursivamente todos os arquivos suportados de uma pasta do Drive."""
    supported_mimes = {
        "application/vnd.google-apps.document",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/pdf",
        "text/plain",
        "text/markdown",
    }

    results = []
    query = f"'{folder_id}' in parents and trashed=false"
    resp = service.files().list(
        q=query, fields="files(id,name,mimeType)", pageSize=200
    ).execute()

    for f in resp.get("files", []):
        if f["mimeType"] == "application/vnd.google-apps.folder":
            sub_name = folder_name + "/" + f["name"] if folder_name else f["name"]
            results += list_drive_files(service, f["id"], sub_name)
        elif f["mimeType"] in supported_mimes:
            results.append({
                "id": f["id"],
                "name": f["name"],
                "mimeType": f["mimeType"],
                "folder": folder_name,
            })

    return results


def download_drive_file(service, file_id: str, mime_type: str) -> bytes | str:
    from googleapiclient.http import MediaIoBaseDownload

    if mime_type == "application/vnd.google-apps.document":
        resp = service.files().export(fileId=file_id, mimeType="text/plain").execute()
        return resp.decode("utf-8") if isinstance(resp, bytes) else resp

    request = service.files().get_media(fileId=file_id)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    buf.seek(0)
    return buf.read()


# ── Extratores de texto ───────────────────────────────────────────────────────

def extract_text(content: bytes | str, mime_type: str, name: str) -> str:
    try:
        if mime_type in ("application/vnd.google-apps.document", "text/plain", "text/markdown"):
            return content if isinstance(content, str) else content.decode("utf-8", errors="ignore")

        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = DocxDocument(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts)

        if mime_type == "application/pdf":
            reader = PdfReader(io.BytesIO(content))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    except Exception as e:
        print(f"  [!] Erro ao extrair texto de {name}: {e}")
    return ""


def read_url(url: str) -> str:
    try:
        headers = {"User-Agent": "Mozilla/5.0 (Mindsight Agent)"}
        r = requests.get(url, headers=headers, timeout=15)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "lxml")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        main = soup.find("main") or soup.find("article") or soup.body
        if not main:
            return ""
        lines = [l.strip() for l in main.get_text(separator="\n").split("\n") if l.strip()]
        return "\n".join(lines)
    except Exception as e:
        print(f"  [!] Erro ao raspar {url}: {e}")
        return ""


# ── Chunking ──────────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> list[str]:
    chunk_size = chunk_size or config.CHUNK_SIZE
    overlap    = overlap    or config.CHUNK_OVERLAP

    if not text.strip():
        return []

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks, current = [], ""

    for para in paragraphs:
        if len(para) > chunk_size:
            if current:
                chunks.append(current.strip())
                current = ""
            for i in range(0, len(para), chunk_size - overlap):
                chunks.append(para[i:i + chunk_size].strip())
            continue

        if len(current) + len(para) + 2 <= chunk_size:
            current = f"{current}\n\n{para}" if current else para
        else:
            if current:
                chunks.append(current.strip())
            tail = chunks[-1][-overlap:] if overlap and chunks else ""
            current = f"{tail}\n\n{para}" if tail else para

    if current.strip():
        chunks.append(current.strip())

    return [c for c in chunks if c.strip()]


# ── Metadados ─────────────────────────────────────────────────────────────────

def extract_category(folder_path: str) -> str:
    """
    Extrai a categoria a partir do nome da primeira subpasta.
    Ex: 'brand/subpasta' → 'brand'
    """
    if not folder_path:
        return "geral"
    parts = folder_path.strip("/").split("/")
    return parts[0].lower() if parts else "geral"


# ── Iterador principal ────────────────────────────────────────────────────────

def iter_documents_from_drive(service) -> Iterator[dict]:
    """Itera por todos os arquivos do Drive e yielda dicts com text + metadata."""
    root_id = config.GOOGLE_DRIVE_FOLDER_ID
    if not root_id:
        print("⚠️  GOOGLE_DRIVE_FOLDER_ID não configurado.")
        return

    files = list_drive_files(service, root_id)
    print(f"📁 {len(files)} arquivo(s) encontrado(s) no Drive\n")

    for f in files:
        print(f"  Lendo: {f['folder']}/{f['name']}")
        try:
            raw = download_drive_file(service, f["id"], f["mimeType"])
            text = extract_text(raw, f["mimeType"], f["name"])
        except Exception as e:
            print(f"  [!] Erro ao baixar {f['name']}: {e}")
            continue

        if not text.strip():
            print(f"  [!] Vazio, pulando: {f['name']}")
            continue

        content_hash = hashlib.md5(text.encode()).hexdigest()
        category = extract_category(f["folder"])

        yield {
            "text": text,
            "metadata": {
                "source": f["name"],
                "drive_id": f["id"],
                "folder": f["folder"],
                "category": category,
                "topic": Path(f["name"]).stem.lower().replace(" ", "-"),
                "content_hash": content_hash,
            }
        }

    # URLs listadas em qualquer arquivo urls.txt no Drive
    for f in files:
        if f["name"].lower() == "urls.txt":
            raw = download_drive_file(service, f["id"], f["mimeType"])
            urls_text = raw if isinstance(raw, str) else raw.decode("utf-8", errors="ignore")
            urls = [l.strip() for l in urls_text.splitlines() if l.strip() and not l.startswith("#")]
            for url in urls:
                print(f"  Raspando: {url}")
                text = read_url(url)
                if not text.strip():
                    continue
                yield {
                    "text": text,
                    "metadata": {
                        "source": url,
                        "drive_id": None,
                        "folder": "urls",
                        "category": "conteudos",
                        "topic": url.split("/")[-1] or "home",
                        "content_hash": hashlib.md5(text.encode()).hexdigest(),
                    }
                }
